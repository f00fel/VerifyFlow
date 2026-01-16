from __future__ import annotations

import re
from functools import lru_cache
from pathlib import Path
from dataclasses import dataclass
from datetime import date
from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple

import yaml
import dateparser

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware

from docx import Document as DocxDocument

import fitz  # PyMuPDF

from natasha import Segmenter, NewsEmbedding, NewsNERTagger, Doc as NatashaDoc
import pymorphy2

app = FastAPI(title="VerifyFlow — VKR checker (DOCX/PDF)")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- NLP init (once) ---
_segmenter = Segmenter()
_emb = NewsEmbedding()
_ner = NewsNERTagger(_emb)
_morph = pymorphy2.MorphAnalyzer()


# ----------------------- Model -----------------------

@dataclass
class Heading:
    text: str
    level: int
    location: str  # "p:12" for docx paragraph index, "page:3" for pdf


@dataclass
class DocumentModel:
    fmt: str               # "docx" | "pdf"
    text: str
    pages: int
    headings: List[Heading]
    meta: Dict[str, Any]


@lru_cache(maxsize=16)
def load_profile(profile: str) -> Dict[str, Any]:
    base_dir = Path(__file__).resolve().parent
    candidates = [
        base_dir / "profiles" / f"{profile}.yaml",
        Path("/app/profiles") / f"{profile}.yaml",
    ]
    for path in candidates:
        if not path.exists():
            continue
        try:
            with path.open("r", encoding="utf-8") as f:
                return yaml.safe_load(f)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to load profile: {e}")
    raise HTTPException(status_code=400, detail=f"Unknown profile: {profile}")


# ----------------------- Extractors -----------------------

def extract_docx(data: bytes) -> DocumentModel:
    doc = DocxDocument(BytesIO(data))

    from collections import Counter

    parts: List[str] = []
    headings: List[Heading] = []

    # --- collect text, headings, and formatting stats in one pass ---
    font_names = []
    font_sizes = []
    line_spacings = []
    paragraphs = doc.paragraphs

    for i, p in enumerate(paragraphs):
        t = (p.text or "").strip()
        if t:
            parts.append(t)

            # heading detection by style name
            style = (p.style.name or "").lower() if p.style is not None else ""
            if "heading" in style or "заголов" in style:
                # try to infer level
                level = 1
                m = re.search(r"(\d+)", style)
                if m:
                    level = max(1, min(6, int(m.group(1))))
                headings.append(Heading(text=t, level=level, location=f"p:{i}"))

        if p.paragraph_format and p.paragraph_format.line_spacing:
            try:
                line_spacings.append(float(p.paragraph_format.line_spacing))
            except Exception:
                pass

        for r in p.runs:
            if r.font.name:
                font_names.append(r.font.name)
            if r.font.size:
                try:
                    font_sizes.append(float(r.font.size.pt))
                except Exception:
                    pass

    most_font = Counter(font_names).most_common(1)
    most_size = Counter(font_sizes).most_common(1)
    most_spacing = Counter(line_spacings).most_common(1)
    if doc.sections:
        sec = doc.sections[0]
        margins_mm = {
            "left": round(mm_from_twips(sec.left_margin.twips), 1),
            "right": round(mm_from_twips(sec.right_margin.twips), 1),
            "top": round(mm_from_twips(sec.top_margin.twips), 1),
            "bottom": round(mm_from_twips(sec.bottom_margin.twips), 1),
        }
    else:
        margins_mm = {}
    text = "\n".join(parts)
    
    # Извлекаем таблицы для проверки календарного плана
    tables = extract_tables_docx(doc)
    
    return DocumentModel(
        fmt="docx",
        text=text,
        pages=1,
        headings=headings,
        meta={
            "paragraphs": len(paragraphs),
            "_docx": doc,
            "tables": tables,
            "detected": {
                "most_common": {
                    "font_name": most_font[0][0] if most_font else None,
                    "font_size": most_size[0][0] if most_size else None,
                    "line_spacing": most_spacing[0][0] if most_spacing else None,
                }
            }
        }
    )


def extract_pdf(data: bytes) -> DocumentModel:
    doc = fitz.open(stream=data, filetype="pdf")
    pages = doc.page_count

    all_lines: List[str] = []
    headings: List[Heading] = []

    # Heuristic: treat as heading if line is mostly uppercase OR looks like a VKR section keyword and is short
    def looks_like_heading(line: str) -> bool:
        s = line.strip()
        if len(s) < 3 or len(s) > 120:
            return False
        letters = re.sub(r"[^A-Za-zА-Яа-яЁё]+", "", s)
        if len(letters) < 3:
            return False
        upper_ratio = sum(1 for c in letters if c.isupper()) / max(1, len(letters))
        if upper_ratio > 0.8:
            return True
        if re.match(r"^(ВВЕДЕНИЕ|ЗАКЛЮЧЕНИЕ|СОДЕРЖАНИЕ|ОГЛАВЛЕНИЕ|СПИСОК|ПРИЛОЖЕНИ)", s.upper()):
            return True
        return False

    for pno in range(pages):
        page = doc.load_page(pno)
        txt = page.get_text("text") or ""
        # normalize hyphenation / line breaks a bit
        txt = re.sub(r"-\n([А-Яа-яЁё])", r"\1", txt)
        lines = [ln.strip() for ln in txt.splitlines() if ln.strip()]
        for ln in lines:
            all_lines.append(ln)
            if looks_like_heading(ln):
                headings.append(Heading(text=ln.strip(), level=1, location=f"page:{pno+1}"))

    text = "\n".join(all_lines)
    
    # Извлекаем таблицы для проверки календарного плана
    tables = extract_tables_pdf(doc)
    
    doc.close()
    return DocumentModel(fmt="pdf", text=text, pages=pages, headings=headings, meta={"pages": pages, "tables": tables})


# ----------------------- Utils -----------------------



def mm_from_twips(twips: int) -> float:
    return (twips / 1440.0) * 25.4

def cm_from_twips(twips: int) -> float:
    return mm_from_twips(twips) / 10.0

def is_black_color(run) -> bool:
    """
    True если цвет чёрный или авто/не задан.
    В Word 'auto' часто значит "чёрный по умолчанию".
    """
    try:
        c = run.font.color
        if c is None:
            return True
        rgb = c.rgb  # can be None
        if rgb is None:
            return True
        return str(rgb).upper() == "000000"
    except Exception:
        return True

def is_bold(run) -> bool:
    try:
        return bool(run.bold) or bool(run.font.bold)
    except Exception:
        return False

def is_italic(run) -> bool:
    try:
        return bool(run.italic) or bool(run.font.italic)
    except Exception:
        return False

def run_font_name(run) -> Optional[str]:
    try:
        return run.font.name
    except Exception:
        return None

def run_font_size_pt(run) -> Optional[float]:
    try:
        if run.font.size is None:
            return None
        return float(run.font.size.pt)
    except Exception:
        return None

def is_latin_text(s: str) -> bool:
    # считаем латиницей, если в строке есть буквы A-Z
    return bool(re.search(r"[A-Za-z]", s or ""))

def norm(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip().upper()


def normalize_fio_key(s: str) -> str:
    """
    Нормализует ФИО/строку для сравнения в игнор-листе: убирает точки и пробелы, делает верхний регистр.
    """
    if not s:
        return ""
    return re.sub(r"[\s\.]+", "", s).upper()


def is_ignored_fullname(name: str, profile: Dict[str, Any]) -> bool:
    """
    Проверяет, содержится ли `name` в списке `ignored_fios` профиля.
    Сравнение выполняется по нормализованным строкам и по фамилии как запасной вариант.
    """
    ignored = profile.get("ignored_fios") or []
    if not ignored:
        return False

    nname = normalize_fio_key(name)
    for ign in ignored:
        iing = normalize_fio_key(ign)
        if not iing:
            continue
        if iing == nname or iing in nname or nname in iing:
            return True
        # запасная проверка по фамилии (последний токен)
        try:
            lname = re.sub(r"[^А-ЯЁ]", "", re.split(r"\s+", name.strip())[-1].upper())
            ilast = re.sub(r"[^А-ЯЁ]", "", re.split(r"\s+", ign.strip())[-1].upper())
            if ilast and lname and ilast == lname:
                return True
        except Exception:
            pass
    return False


def is_likely_not_date(text: str, match_start: int, match_end: int) -> bool:
    """
    Проверяет, не является ли найденный паттерн номером специальности или документа.
    Возвращает True, если это НЕ дата (т.е. это номер специальности/документа).
    """
    # Расширяем контекст вокруг совпадения
    context_start = max(0, match_start - 50)
    context_end = min(len(text), match_end + 50)
    context = text[context_start:context_end].lower()
    
    # Исключаем номера специальностей (например, "09.03.04 Программная инженерия")
    if re.search(r"(направление|специальность|код|профиль).*?\d{2}\.\d{2}\.\d{2}", context):
        return True
    
    # Исключаем номера документов (например, "№ 33.02-05/334")
    if re.search(r"№\s*\d+\.\d+[-/]\d+", context):
        return True
    
    # Исключаем паттерны типа "XX.XX.XX" без "г." после них (через пробел) и без контекста даты
    # Проверяем, нет ли "г." после паттерна через пробел
    text_after_match = text[match_end:match_end+10]
    has_g_after = bool(re.search(r"\s+г\.", text_after_match))
    
    # Если нет "г." после паттерна, проверяем контекст на специальность
    if not has_g_after:
        after_match = text[match_end:match_end+100].lower()
        if re.search(r"(программная|инженерия|направление|специальность|код|профиль)", after_match):
            return True
    
    return False


def parse_dates_with_context(text: str) -> List[Tuple[date, str, int]]:
    """
    Находит даты в форматах:
    - "09 декабря 2024 г." (словесная дата с "г.")
    - "06.12.2024 г." (цифровая дата с "г.")
    - "06.12.2024" (цифровая дата без "г.", но в контексте даты)
    
    Исключает номера специальностей (09.03.04) и документов (33.02-05).
    Returns list of (date_obj, raw_fragment, index_in_text)
    """
    res: List[Tuple[date, str, int]] = []
    
    # 1. Словесные даты: "09 декабря 2024 г." или "9 декабря 2024 г."
    months_ru = [
        "января", "февраля", "марта", "апреля", "мая", "июня",
        "июля", "августа", "сентября", "октября", "ноября", "декабря"
    ]
    months_pattern = "|".join(months_ru)
    
    # Паттерн для словесной даты с "г."
    word_date_pattern = r"\b(\d{1,2})\s+(" + months_pattern + r")\s+(\d{4})\s+г\.\b"
    for m in re.finditer(word_date_pattern, text, re.IGNORECASE):
        day_str = m.group(1)
        month_name = m.group(2).lower()
        year_str = m.group(3)
        
        try:
            month_num = months_ru.index(month_name) + 1
            day = int(day_str)
            year = int(year_str)
            
            if 1 <= day <= 31 and 1900 <= year <= 2100:
                dt = date(year, month_num, day)
                frag = m.group(0)
                res.append((dt, frag, m.start()))
        except (ValueError, IndexError):
            continue
    
    # 2. Цифровые даты с "г.": "06.12.2024 г." или "6.12.2024 г."
    date_with_g_pattern = r"\b(\d{1,2})\.(\d{1,2})\.(\d{4})(?:\s+г\.)?\b"
    for m in re.finditer(date_with_g_pattern, text):
        day_str, month_str, year_str = m.groups()
        
        # Проверяем, не является ли это номером специальности/документа
        if is_likely_not_date(text, m.start(), m.end()):
            continue
        
        try:
            day = int(day_str)
            month = int(month_str)
            year = int(year_str)
            
            if 1 <= day <= 31 and 1 <= month <= 12 and 1900 <= year <= 2100:
                dt = date(year, month, day)
                frag = m.group(0)
                res.append((dt, frag, m.start()))
        except (ValueError, OverflowError):
            continue
    
    # 3. Цифровые даты без "г.", но в контексте, указывающем на дату
    # Ищем даты в формате ДД.ММ.ГГГГ, которые находятся рядом с маркерами даты
    date_markers = ["дата", "число", "год", "принял", "заверш", "утвержден", "подписан"]
    date_pattern = r"\b(\d{1,2})\.(\d{1,2})\.(\d{4})\b"
    
    for m in re.finditer(date_pattern, text):
        # Проверяем контекст вокруг даты
        context_start = max(0, m.start() - 30)
        context_end = min(len(text), m.end() + 10)
        context = text[context_start:context_end].lower()
        
        # Если есть маркер даты или "г." после даты (через пробел)
        has_date_marker = any(marker in context for marker in date_markers)
        # Проверяем наличие "г." после даты через пробел (до 10 символов после даты)
        text_after = text[m.end():m.end()+10]
        has_g_after = bool(re.search(r"\s+г\.", text_after))
        
        # Пропускаем, если это похоже на номер специальности/документа
        if is_likely_not_date(text, m.start(), m.end()):
            continue
        
        # Принимаем только если есть маркер даты или "г." после (через пробел)
        if has_date_marker or has_g_after:
            day_str, month_str, year_str = m.groups()
            try:
                day = int(day_str)
                month = int(month_str)
                year = int(year_str)
                
                if 1 <= day <= 31 and 1 <= month <= 12 and 1900 <= year <= 2100:
                    dt = date(year, month, day)
                    frag = m.group(0)
                    res.append((dt, frag, m.start()))
            except (ValueError, OverflowError):
                continue
    
    # 4. Только год с "г.": "2024 г."
    year_only_pattern = r"\b(20\d{2})\s+г\.\b"
    for m in re.finditer(year_only_pattern, text):
        year_str = m.group(1)
        try:
            year = int(year_str)
            if 1900 <= year <= 2100:
                dt = date(year, 1, 1)
                frag = m.group(0)
                res.append((dt, frag, m.start()))
        except ValueError:
            continue
    
    # de-dup by position
    seen = set()
    uniq = []
    for d, frag, idx in sorted(res, key=lambda x: x[2]):
        key = (idx, frag)
        if key in seen:
            continue
        seen.add(key)
        uniq.append((d, frag, idx))
    return uniq


def extract_person_names(text: str) -> List[Tuple[str, int]]:
    """
    Use Natasha NER to find PER entities, return (name_text, start_index).
    """
    doc = NatashaDoc(text)
    doc.segment(_segmenter)
    doc.tag_ner(_ner)
    out: List[Tuple[str, int]] = []
    for span in doc.spans:
        if span.type == "PER":
            out.append((span.text, span.start))
    return out


def extract_name_fallback(text: str) -> List[Tuple[str, int]]:
    """
    Простая эвристика: ищет последовательности из 2-3 слов с заглавной буквы (Ф И О).
    Возвращает список (name, start_index).
    """
    res: List[Tuple[str, int]] = []
    # Ищем 2-3 слова с заглавной кириллической буквы
    for m in re.finditer(r"\b[А-ЯЁ][а-яё]+(?:[-\s][А-ЯЁ][а-яё]+)?\s+[А-ЯЁ][а-яё]+(?:\s+[А-ЯЁ][а-яё]+)?\b", text):
        res.append((m.group(0), m.start()))
    return res


def find_next_capitalized_sequence(text: str, start: int) -> Optional[int]:
    """Find position of next sequence of 2-3 capitalized Cyrillic words after `start`.
    Returns absolute index or None.
    """
    sub = text[start:]
    m = re.search(r"[А-ЯЁ][а-яё]+(?:\s+[А-ЯЁ][а-яё]+){1,2}", sub)
    if m:
        return start + m.start()
    return None


def guess_case_of_fullname(name: str) -> Optional[str]:
    """
    Very rough: determine grammatical case by first token that morph can parse.
    Returns pymorphy2 case like 'nomn', 'gent', 'datv', ...
    """
    tokens = [t for t in re.split(r"\s+", name.strip()) if t]
    for t in tokens:
        w = re.sub(r"[^А-Яа-яЁё-]", "", t)
        if not w:
            continue
        parses = _morph.parse(w)
        if not parses:
            continue
        # take best parse
        tag = parses[0].tag
        if tag.case:
            return tag.case
    return None


def case_name_ru(code: Optional[str]) -> str:
    """Возвращает название падежа на русском по коду pymorphy2 (nomn, gent, ...)."""
    if not code:
        return ""
    mapping = {
        "nomn": "именительный",
        "gent": "родительный",
        "datv": "дательный",
        "accs": "винительный",
        "ablt": "творительный",
        "loct": "предложный",
        "voct": "звательный",
    }
    return mapping.get(code, code)


def has_abbreviated_name(name: str) -> bool:
    """
    Проверяет, содержит ли ФИО сокращения типа И.О. или И. О.
    """
    # Паттерны: И.О., И. О., И.О, И О. и т.п.
    if re.search(r"\b[А-ЯЁ]\.\s*[А-ЯЁ]\.?\b", name):
        return True
    # Проверка на одиночные инициалы с точкой
    parts = name.split()
    for part in parts:
        if len(part) == 2 and part.endswith('.') and part[0].isupper():
            return True
    return False


def find_person_after_anchors(text: str, anchors: List[str], window_after: int = 400) -> List[Tuple[str, str, int]]:
    """
    Ищет персоны после списка якорных фраз (anchors).
    Возвращает список кортежей (name, anchor_used, absolute_index).
    Использует Natasha NER, затем fallback-эвристику.
    """
    res: List[Tuple[str, str, int]] = []
    text_low = text.lower()
    for anchor in anchors:
        a_low = anchor.lower()
        start = 0
        while True:
            idx = text_low.find(a_low, start)
            if idx == -1:
                break
            # окно после якоря
            start_search = idx + len(a_low)
            window = text[start_search:start_search + window_after]
            # сначала NER
            persons = extract_person_names(window)
            if not persons:
                persons = extract_name_fallback(window)
            if persons:
                name, relpos = persons[0]
                res.append((name, anchor, start_search + relpos))
            start = idx + 1
    return res


def _anchor_conflict_resolver(text: str, anchor: str) -> Optional[str]:
    """
    Решает конфликт: если после 'руководитель' идёт 'допустить' или 'обучающегося',
    то возвращаем альтернативный якорь 'допустить' чтобы обработать как студент.
    Иначе возвращает None.
    """
    if 'руководитель' not in anchor.lower():
        return None
    # смотрим небольшой кусок после слова 'руководитель'
    pos = text.lower().find(anchor.lower())
    if pos == -1:
        return None
    tail = text[pos + len(anchor): pos + len(anchor) + 60].lower()
    if 'допустить' in tail or 'обучающ' in tail:
        return 'допустить'
    return None


def parse_strict_date_format(text: str) -> List[Tuple[date, str, int]]:
    """
    Парсит даты в форматах:
    - "09 декабря 2024 г." или "09 декабря 2024"
    - "06.12.2024 г." или "06.12.2024"

    Исключает номера специальностей и документов.
    Возвращает список (date_obj, raw_fragment, index_in_text).
    """
    res: List[Tuple[date, str, int]] = []

    # 1. Словесные даты: "09 декабря 2024" (+/- "г.")
    months_ru = [
        "января", "февраля", "марта", "апреля", "мая", "июня",
        "июля", "августа", "сентября", "октября", "ноября", "декабря"
    ]
    months_pattern = "|".join(months_ru)

    word_date_pattern = r"\b(\d{1,2})\s+(" + months_pattern + r")\s+(\d{4})(?:[\u00A0\s]*г\.)?\b"
    for m in re.finditer(word_date_pattern, text, re.IGNORECASE):
        day_str = m.group(1)
        month_name = m.group(2).lower()
        year_str = m.group(3)

        try:
            month_num = months_ru.index(month_name) + 1
            day = int(day_str)
            year = int(year_str)

            if 1 <= day <= 31 and 1900 <= year <= 2100:
                dt = date(year, month_num, day)
                frag = m.group(0)
                res.append((dt, frag, m.start()))
        except (ValueError, IndexError):
            continue

    # 2. Цифровые даты: "06.12.2024" (+/- "г.")
    date_with_g_pattern = r"\b(\d{1,2})\.(\d{1,2})\.(\d{4})(?:[\u00A0\s]*г\.)?\b"
    for m in re.finditer(date_with_g_pattern, text):
        # Отсекаем то, что по контексту может быть не датой (номер документа)
        if is_likely_not_date(text, m.start(), m.end()):
            continue

        day_str, month_str, year_str = m.groups()
        try:
            day = int(day_str)
            month = int(month_str)
            year = int(year_str)

            if 1 <= day <= 31 and 1 <= month <= 12 and 1900 <= year <= 2100:
                dt = date(year, month, day)
                frag = m.group(0)
                res.append((dt, frag, m.start()))
        except (ValueError, OverflowError):
            continue

    return res


def find_signature_fields(text: str) -> Dict[str, bool]:
    """
    Ищет поля для подписей студента и руководителя.
    Возвращает словарь с найденными полями.
    """
    text_low = text.lower()
    found = {
        "student_signature": False,
        "supervisor_signature": False,
    }
    
    # Паттерны для подписи студента
    student_patterns = [
        r"подпись\s+студент",
        r"подпись\s+обучающ",
        r"студент\s*[:\s]*\s*подпись",
        r"обучающ[ийея]\s*[:\s]*\s*подпись",
    ]
    
    # Паттерны для подписи руководителя
    supervisor_patterns = [
        r"подпись\s+руководител",
        r"руководител[ья]\s*[:\s]*\s*подпись",
        r"научный\s+руководитель\s*[:\s]*\s*подпись",
    ]
    
    for pattern in student_patterns:
        if re.search(pattern, text_low):
            found["student_signature"] = True
            break
    
    for pattern in supervisor_patterns:
        if re.search(pattern, text_low):
            found["supervisor_signature"] = True
            break
    
    return found


def check_text_formatting(text: str) -> List[Tuple[str, int]]:
    """
    Проверяет форматирование текста на лишние пробелы и нестандартные символы.
    Возвращает список (проблема, позиция).
    """
    issues: List[Tuple[str, int]] = []

    # Оставляем только детекцию нестандартных символов (неразрывные пробелы, мягкие переносы и пр.).
    for m in re.finditer(r"[\u00A0\u00AD\u2000-\u200F\u2028-\u202F]", text):
        issues.append(("nonstandard_char", m.start()))

    return issues


def extract_tables_docx(doc) -> List[Dict[str, Any]]:
    """
    Извлекает таблицы из DOCX документа.
    Возвращает список словарей с данными таблиц.
    """
    tables_data = []
    for table_idx, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            if any(cell_text for cell_text in row_data):  # Пропускаем пустые строки
                rows_data.append(row_data)
        if rows_data:
            tables_data.append({
                "index": table_idx,
                "rows": rows_data,
                "row_count": len(rows_data),
                "col_count": len(rows_data[0]) if rows_data else 0,
            })
    return tables_data


def extract_tables_pdf(doc) -> List[Dict[str, Any]]:
    """
    Извлекает таблицы из PDF документа (базовая реализация через текст).
    """
    tables_data = []
    # PyMuPDF имеет ограниченную поддержку таблиц, используем эвристику
    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)
        text = page.get_text("text")
        # Ищем структурированные данные, похожие на таблицы (много табуляций или выравнивание)
        lines = text.split('\n')
        potential_table_rows = []
        for line in lines:
            # Если строка содержит много разделителей (табуляция, множественные пробелы)
            if re.search(r'\t| {3,}', line):
                cells = re.split(r'\t| {3,}', line)
                cells = [c.strip() for c in cells if c.strip()]
                if len(cells) >= 2:
                    potential_table_rows.append(cells)
        
        if len(potential_table_rows) >= 2:  # Минимум 2 строки для таблицы
            tables_data.append({
                "index": len(tables_data),
                "rows": potential_table_rows,
                "row_count": len(potential_table_rows),
                "col_count": max(len(row) for row in potential_table_rows) if potential_table_rows else 0,
                "page": page_num + 1,
            })
    
    return tables_data


# ----------------------- Rules -----------------------

def rule_margins_docx(dm: DocumentModel, profile: Dict[str, Any]) -> List[Dict[str, Any]]:
    cfg = profile.get("margins_docx") or {}
    if not cfg.get("enabled", False):
        return []
    if dm.fmt != "docx":
        return []

    sev = cfg.get("severity", "critical")
    exp = cfg.get("expected_mm") or {}
    tol = float(cfg.get("tolerance_mm", 1.5))

    doc = dm.meta.get("_docx")
    if doc is None or not getattr(doc, "sections", None):
        return []

    sec = doc.sections[0]
    margins = {
        "left": mm_from_twips(sec.left_margin.twips),
        "right": mm_from_twips(sec.right_margin.twips),
        "top": mm_from_twips(sec.top_margin.twips),
        "bottom": mm_from_twips(sec.bottom_margin.twips),
    }

    issues = []
    for side in ("left", "right", "top", "bottom"):
        if side not in exp:
            continue
        if abs(margins[side] - float(exp[side])) > tol:
            issues.append({
                "severity": sev,
                "type": "formal",
                "rule": f"Margins.{side}",
                "message": f"Поле {side}: {margins[side]:.1f} мм (ожидается {float(exp[side]):.0f} мм)",
                "evidence": f"{side}={margins[side]:.1f}мм",
                "location": "docx:section1",
                "how_to_fix": "Word → Разметка страницы → Поля → Настраиваемые поля.",
            })
    return issues

def rule_indent_docx(dm: DocumentModel, profile: Dict[str, Any]) -> List[Dict[str, Any]]:
    cfg = profile.get("indent_docx") or {}
    if not cfg.get("enabled", False):
        return []
    if dm.fmt != "docx":
        return []
    # Правило отключено — возвращаем пустой список (пользователь запросил удаление проверки).
    return []

def rule_title_fields(dm: DocumentModel, profile: Dict[str, Any]) -> List[Dict[str, Any]]:
    cfg = (profile.get("title_page_fields") or {})
    if not cfg.get("enabled", True):
        return []
    sev = cfg.get("severity", "warning")
    title_pages = int(cfg.get("title_scope_pages", 1))
    required = cfg.get("required") or []

    # Determine "title scope"
    if dm.fmt == "pdf":
        # crude: use first N pages worth of text by splitting approx
        # We'll take first ~5000 chars per page as an estimate
        scope = dm.text[: title_pages * 5000]
        loc = f"pages:1..{title_pages}"
    else:
        scope = dm.text[: 6000]
        loc = "start of document"

    scope_low = scope.lower()

    issues: List[Dict[str, Any]] = []

    # year check separately
    year_found = re.search(r"\b(20\d{2})\b", scope)
    if not year_found:
        issues.append({
            "severity": sev,
            "type": "formal",
            "rule": "Title.Year",
            "message": "На титульной части не найден год (например, 2026).",
            "evidence": "—",
            "location": loc,
            "how_to_fix": "Проверь титульный лист: обычно внизу указывают город и год.",
        })

    for field in required:
        key = field.get("key", "field")
        hints = [h.lower() for h in (field.get("hints") or [])]

        found_anchor = any(h in scope_low for h in hints)

        if key == "student_fio" and found_anchor:
            # дополнительно проверяем, есть ли ФИО рядом
            persons = extract_person_names(scope)
            if persons:
                continue  # ФИО найдено → всё ок

        if not found_anchor:
            issues.append({
                "severity": sev,
                "type": "formal",
                "rule": "Title.RequiredField",
                "message": f"В титульной зоне не найдено поле/маркер: {key}",
                "evidence": ", ".join(hints[:4]) if hints else key,
                "location": loc,
                "how_to_fix": "Проверь титульный лист: должны быть ФИО, тема, руководитель, город и год (согласно методичке).",
            })

    return issues


def rule_dates(dm: DocumentModel, profile: Dict[str, Any]) -> List[Dict[str, Any]]:
    cfg = (profile.get("dates") or {})
    if not cfg.get("enabled", True):
        return []
    sev = cfg.get("severity", "warning")

    dates = parse_dates_with_context(dm.text)
    issues: List[Dict[str, Any]] = []

    if len(dates) == 0:
        issues.append({
            "severity": sev,
            "type": "formal",
            "rule": "Dates.None",
            "message": "В документе не найдено дат (форматы: '06.12.2024 г.' или '09 декабря 2024 г.').",
            "evidence": "—",
            "location": "document",
            "how_to_fix": "Если документ служебный (задание/утверждение), добавь дату в требуемом формате с 'г.' после даты.",
        })
        return issues

    # Check suspicious years
    years = [d.year for d, _, _ in dates]
    if years:
        miny, maxy = min(years), max(years)
        if maxy - miny > 15:
            issues.append({
                "severity": sev,
                "type": "logical",
                "rule": "Dates.Range",
                "message": f"Подозрительно широкий разброс годов: {miny}–{maxy}.",
                "evidence": f"{miny}..{maxy}",
                "location": "document",
                "how_to_fix": "Проверь корректность дат: нет ли опечаток в годах (например, 2014 вместо 2024).",
            })

    # Try detect date ranges (start/end) by nearby words
    range_cfg = cfg.get("range_hints") or {}
    start_hints = [h.lower() for h in range_cfg.get("start", [])]
    end_hints = [h.lower() for h in range_cfg.get("end", [])]

    # Very simple: find two nearest dates with 'с'/'по' or 'от'/'до' between
    # We'll scan text windows around each date
    pairs: List[Tuple[date, date, str]] = []
    for i in range(len(dates)-1):
        d1, f1, idx1 = dates[i]
        d2, f2, idx2 = dates[i+1]
        between = dm.text[idx1:idx2].lower()
        if any(h in between for h in ["с", "от"] + start_hints) and any(h in between for h in ["по", "до"] + end_hints):
            pairs.append((d1, d2, f"{f1} … {f2}"))

    for d1, d2, frag in pairs[:5]:
        if d1 > d2:
            issues.append({
                "severity": sev,
                "type": "logical",
                "rule": "Dates.Order",
                "message": "Нарушена последовательность дат в диапазоне (начало позже окончания).",
                "evidence": frag,
                "location": "document",
                "how_to_fix": "Проверь, что дата начала меньше или равна дате окончания.",
            })

    return issues


def rule_fio_cases(dm: DocumentModel, profile: Dict[str, Any]) -> List[Dict[str, Any]]:
    cfg = (profile.get("fio_cases") or {})
    if not cfg.get("enabled", True):
        return []
    sev = cfg.get("severity", "warning")
    contexts = cfg.get("contexts") or []

    issues: List[Dict[str, Any]] = []

    # We'll search within limited windows after triggers and extract PER entities there
    text = dm.text
    text_low = text.lower()

    for ctx in contexts:
        trigger = (ctx.get("trigger") or "").lower()
        expected = ctx.get("expected_case") or None
        label = ctx.get("label") or trigger

        if not trigger or not expected:
            continue
        
        # Пропускаем контексты с "руководитель" — они обрабатываются в rule_supervisor_fio_detailed
        if "руководитель" in trigger:
            continue

        for m in re.finditer(re.escape(trigger), text_low):
            start = m.end()
            window = text[start:start+140]  # after trigger
            persons = extract_person_names(window)
            if not persons:
                continue
            # take first person in window
            name, pos = persons[0]
            # Пропускаем, если имя в списке игнорируемых
            if is_ignored_fullname(name, profile):
                continue
            case = guess_case_of_fullname(name)
            # Accept genitive or dative when expected is genitive
            if case and case != expected:
                if not (expected == "gent" and case == "datv"):
                    found_ru = case_name_ru(case)
                    expected_ru = case_name_ru(expected)
                    issues.append({
                        "severity": sev,
                        "type": "syntactic",
                        "rule": "FIO.Case",
                        "message": f"ФИО после контекста “{label}” возможно в неверном падеже: найден {found_ru}, ожидается {expected_ru}.",
                        "evidence": f"{trigger} … {name}",
                        "location": "document",
                        "how_to_fix": "Исправь ФИО в формальном поле в соответствии с требуемым падежом (обычно 'от студента' → родительный).",
                    })

    return issues


def rule_student_fio_detailed(dm: DocumentModel, profile: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Детальная проверка ФИО обучающегося:
    - Родительный падеж
    - Отсутствие сокращений (И.О.)
    - Орфография (базовая проверка через морфологию)
    """
    cfg = profile.get("student_fio_detailed") or {}
    if not cfg.get("enabled", False):
        return []

    sev = cfg.get("severity", "critical")
    issues: List[Dict[str, Any]] = []

    # Сфокусируем поиск на титульной зоне
    title_scope = dm.text[:6000] if len(dm.text) > 6000 else dm.text

    # По заданию — используем якоря 'обучающегося' и 'допустить' как основные
    anchors = ["обучающегося", "допустить"]

    found = find_person_after_anchors(title_scope, anchors, window_after=300)

    if not found:
        # Если ничего не найдено — оформляем ошибку
        issues.append({
            "severity": sev,
            "type": "formal",
            "rule": "StudentFIO.Missing",
            "message": "Не найдено ФИО обучающегося в титульной зоне (якоря: обучающегося/допустить).",
            "evidence": "—",
            "location": "title_page",
            "how_to_fix": "Укажи полное ФИО обучающегося после маркера 'обучающегося' или 'допустить'.",
        })
        return issues

    # Обрабатываем найденные персоны (берём первую релевантную)
    for name, anchor, abs_idx in found:
        if is_ignored_fullname(name, profile):
            continue

        # Проверяем на сокращения — это всегда ошибка
        if has_abbreviated_name(name):
            issues.append({
                "severity": sev,
                "type": "syntactic",
                "rule": "StudentFIO.Abbreviation",
                "message": "ФИО обучающегося содержит сокращения (И.О.). Требуется полное написание (Фамилия Имя Отчество).",
                "evidence": f"{anchor} ... {name}",
                "location": "title_page",
                "how_to_fix": "Укажи полное ФИО без сокращений (например, 'Иванов Иван Иванович').",
            })

        # Проверка падежа: допускаем родительный, дательный и винительный
        case = guess_case_of_fullname(name)
        accepted_cases = ("gent", "datv", "accs")

        if case and case not in accepted_cases:
            found_ru = case_name_ru(case)
            accepted_ru = ", ".join(case_name_ru(c) for c in accepted_cases)
            issues.append({
                "severity": sev,
                "type": "syntactic",
                "rule": "StudentFIO.Case",
                "message": f"ФИО обучающегося должно быть полным в одном из падежей: родительном, дательном или винительном. Найден: {found_ru}.",
                "evidence": f"{anchor} ... {name}",
                "location": "title_page",
                "how_to_fix": "Укажи полное ФИО в одном из допустимых падежей (например, 'от студента Иванова Ивана Ивановича').",
            })

    return issues

def rule_supervisor_fio_detailed(dm: DocumentModel, profile: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Детальная проверка ФИО руководителя:
    - Косвенный падеж (обычно именительный или родительный)
    - Наличие учёной степени, звания, должности
    """
    cfg = profile.get("supervisor_fio_detailed") or {}
    if not cfg.get("enabled", False):
        return []

    sev = cfg.get("severity", "critical")
    issues: List[Dict[str, Any]] = []

    title_scope = dm.text[:6000] if len(dm.text) > 6000 else dm.text

    # Якоря для руководителя (включая возможные варианты с пробелами/точками)
    anchors = ["руководитель", "научный руководитель", "руководителя", "руководитель .", "руководитель .\t", "руководитель .\n"]

    # Ищем сначала кандидатов по 'руководитель'
    candidates = find_person_after_anchors(title_scope, anchors, window_after=400)

    # Разрешаем конфликты: если после 'руководитель' идёт 'допустить' или 'обучающегося',
    # то передаём обработку этому якорю (т.е. считаем, что это студент)
    filtered: List[Tuple[str, str, int]] = []
    for name, anchor, idx in candidates:
        alt = _anchor_conflict_resolver(title_scope, anchor)
        if alt:
            # добавим в профиль как не найдено для руководителя — студенческая логика должна поймать это
            continue
        filtered.append((name, anchor, idx))

    if not filtered:
        issues.append({
            "severity": sev,
            "type": "formal",
            "rule": "SupervisorFIO.Missing",
            "message": "Не найдено ФИО руководителя в титульной зоне (якорь: руководитель).",
            "evidence": "—",
            "location": "title_page",
            "how_to_fix": "Укажи полное ФИО руководителя после маркера 'руководитель'.",
        })
        return issues

    # Берём первого релевантного кандидата
    name, anchor, abs_idx = filtered[0]
    if is_ignored_fullname(name, profile):
        return []

    # Проверка на сокращения — недопустимо
    if has_abbreviated_name(name):
        issues.append({
            "severity": sev,
            "type": "syntactic",
            "rule": "SupervisorFIO.Abbreviation",
            "message": "ФИО руководителя содержит сокращения. Требуется полное написание (Фамилия Имя Отчество).",
            "evidence": f"{anchor} ... {name}",
            "location": "title_page",
            "how_to_fix": "Укажи полное ФИО руководителя без сокращений.",
        })

    # Проверка падежа: требуется косвенный падеж (не именительный).
    case = guess_case_of_fullname(name)
    if case == 'nomn':
        issues.append({
            "severity": sev,
            "type": "syntactic",
            "rule": "SupervisorFIO.Case",
            "message": "ФИО руководителя должно быть полным и в косвенном падеже (не именительный).",
            "evidence": f"{anchor} ... {name}",
            "location": "title_page",
            "how_to_fix": "Укажи ФИО руководителя в косвенном падеже (например, 'руководитель Иванова Ивана Ивановича').",
        })

    return issues

def rule_date_format_strict(dm: DocumentModel, profile: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Требует корректный формат дат: "ДД.ММ.ГГГГ" (+/- "г.") или "ДД месяц ГГГГ" (+/- "г.").
    Дополнительно ловит битые цифровые даты (неправильный порядок или выход за диапазон).
    """
    cfg = profile.get("date_format_strict") or {}
    if not cfg.get("enabled", False):
        return []

    sev = cfg.get("severity", "warning")
    issues: List[Dict[str, Any]] = []

    strict_dates = parse_strict_date_format(dm.text)
    all_dates = parse_dates_with_context(dm.text)
    strict_positions = {idx: frag for _, frag, idx in strict_dates}

    for dt, frag, idx in all_dates:
        if idx not in strict_positions and not is_likely_not_date(dm.text, idx, idx + len(frag)):
            issues.append({
                "severity": sev,
                "type": "formal",
                "rule": "Date.Format",
                "message": "Дата должна быть в формате 'ДД.ММ.ГГГГ' или 'ДД месяц ГГГГ' (допустимо с/без 'г.').",
                "evidence": frag,
                "location": "document",
                "how_to_fix": "Приведи дату к виду '06.12.2024' или '06.12.2024 г.' / '09 декабря 2024 г.'.",
            })

    # Плохие цифровые даты (перепутан порядок или выход за диапазон)
    loose_pattern = re.compile(r"\b(\d{1,4})\.(\d{1,3})\.(\d{2,4})(?:[\u00A0\s]*г\.)?\b")
    invalid_seen = set()
    for m in loose_pattern.finditer(dm.text):
        day_str, month_str, year_str = m.groups()
        frag = m.group(0)
        idx = m.start()
        if idx in strict_positions or idx in invalid_seen:
            continue

        # Игнорируем коды направлений/специальностей вида 09.03.01 и подобные
        context = dm.text[max(0, idx - 30): idx + 40].lower()
        markers = ["направлен", "специальн", "профиль", "информат", "техника", "инженер", "программы"]
        if len(year_str) == 2 and any(mk in context for mk in markers):
            continue

        try:
            day = int(day_str)
            month = int(month_str)
            year = int(year_str)
        except ValueError:
            invalid = True
        else:
            invalid = (
                len(day_str) > 2 or len(month_str) > 2 or len(year_str) != 4 or
                not (1 <= day <= 31 and 1 <= month <= 12 and 1900 <= year <= 2100)
            )
        if invalid:
            invalid_seen.add(idx)
            issues.append({
                "severity": sev,
                "type": "formal",
                "rule": "Date.InvalidFormat",
                "message": "Некорректная дата. Ожидается порядок ДД.ММ.ГГГГ (допустимо с/без 'г.').",
                "evidence": frag,
                "location": "document",
                "how_to_fix": "Проверь порядок дня/месяца и значения (1-31 и 1-12), год — четыре цифры.",
            })

    return issues

def rule_signatures(dm: DocumentModel, profile: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Проверка наличия полей для подписей студента и руководителя.
    """
    # Проверка подписей отключена по запросу — возвращаем пустой список проблем.
    return []

def rule_text_formatting(dm: DocumentModel, profile: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Проверка форматирования текста: лишние пробелы, нестандартные символы.
    """
    cfg = profile.get("text_formatting") or {}
    if not cfg.get("enabled", False):
        return []
    
    sev = cfg.get("severity", "warning")
    issues: List[Dict[str, Any]] = []
    
    formatting_issues = check_text_formatting(dm.text)
    
    if not formatting_issues:
        return []
    
    # Группируем по типам проблем
    by_type: Dict[str, List[int]] = {}
    for issue_type, pos in formatting_issues[:50]:  # Ограничиваем количество
        if issue_type not in by_type:
            by_type[issue_type] = []
        by_type[issue_type].append(pos)
    
    type_labels = {
        "double_space": "Двойные пробелы",
        "space_before_punct": "Пробелы перед знаками препинания",
        "nonstandard_char": "Нестандартные символы",
        "tab_char": "Табуляции в тексте",
    }
    
    for issue_type, positions in by_type.items():
        count = len(positions)
        label = type_labels.get(issue_type, issue_type)
        issues.append({
            "severity": sev,
            "type": "formal",
            "rule": f"TextFormatting.{issue_type}",
            "message": f"Найдено проблем форматирования типа '{label}': {count}",
            "evidence": f"Первые позиции: {', '.join(map(str, positions[:5]))}",
            "location": "document",
            "how_to_fix": "Исправь форматирование: убери лишние пробелы, замени табуляции на пробелы, проверь символы.",
        })
    
    return issues

def rule_calendar_plan(dm: DocumentModel, profile: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Проверка календарного плана: структура таблицы, корректность дат, хронологический порядок.
    """
    cfg = profile.get("calendar_plan") or {}
    if not cfg.get("enabled", False):
        return []
    
    sev = cfg.get("severity", "critical")
    issues: List[Dict[str, Any]] = []
    
    # Извлекаем таблицы
    tables = []
    if dm.fmt == "docx":
        doc = dm.meta.get("_docx")
        if doc:
            tables = extract_tables_docx(doc)
    elif dm.fmt == "pdf":
        # Для PDF нужен доступ к документу, сохраняем в meta при извлечении
        tables = dm.meta.get("tables", [])
    
    if not tables:
        issues.append({
            "severity": "warning",
            "type": "structural",
            "rule": "CalendarPlan.NotFound",
            "message": "Не найдена таблица календарного плана.",
            "evidence": "—",
            "location": "document",
            "how_to_fix": "Добавь таблицу календарного плана с этапами работы и датами.",
        })
        return issues
    
    # Ищем таблицу, похожую на календарный план (содержит даты и этапы)
    calendar_table = None
    loose_pattern = re.compile(r"\b(\d{1,4})\.(\d{1,3})\.(\d{2,4})(?:[\u00A0\s]*г\.)?\b")
    for table in tables:
        rows = table.get("rows", [])
        if len(rows) < 2:
            continue
        
        # Проверяем, есть ли в таблице даты (с "г." или словесные)
        has_dates = False
        for row in rows:
            row_text = " ".join(row)
            # Проверяем наличие дат с/без "г." или словесных дат
            if parse_strict_date_format(row_text) or loose_pattern.search(row_text) or re.search(r"\d{1,2}\s+(января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)\s+\d{4}(?:\s+г\.)?", row_text, re.IGNORECASE):
                has_dates = True
                break
        
        if has_dates:
            calendar_table = table
            break
    
    if not calendar_table:
        issues.append({
            "severity": "warning",
            "type": "structural",
            "rule": "CalendarPlan.NoDates",
            "message": "В таблицах не найдены даты в формате 'ДД.ММ.ГГГГ г.' или 'ДД месяц ГГГГ г.'.",
            "evidence": "—",
            "location": "document",
            "how_to_fix": "Проверь, что календарный план содержит даты в правильном формате с 'г.' после даты (например, '06.12.2024 г.' или '09 декабря 2024 г.').",
        })
        return issues
    
    # Проверяем хронологический порядок дат
    dates_found = []
    rows = calendar_table.get("rows", [])
    
    for row_idx, row in enumerate(rows[1:], start=1):  # Пропускаем заголовок
        row_text = " ".join(row)
        dates_in_row = parse_strict_date_format(row_text)
        strict_spans = [(pos, pos + len(frag)) for _, frag, pos in dates_in_row]
        for dt, frag, pos in dates_in_row:
            dates_found.append((dt, frag, row_idx))

        # Ловим явные ошибки формата в строках таблицы (даже если даты не распарсились)
        for m in loose_pattern.finditer(row_text):
            frag = m.group(0)
            start, end = m.start(), m.end()
            if any(s <= start < e for s, e in strict_spans):
                continue
            try:
                day = int(m.group(1))
                month = int(m.group(2))
                year = int(m.group(3))
            except ValueError:
                invalid = True
            else:
                invalid = (
                    len(m.group(1)) > 2 or len(m.group(2)) > 2 or len(m.group(3)) != 4 or
                    not (1 <= day <= 31 and 1 <= month <= 12 and 1900 <= year <= 2100)
                )
            if invalid:
                issues.append({
                    "severity": sev,
                    "type": "formal",
                    "rule": "CalendarPlan.InvalidDate",
                    "message": "Некорректная дата в календарном плане. Ожидается формат ДД.ММ.ГГГГ (допустимо с/без 'г.').",
                    "evidence": frag,
                    "location": "calendar_plan_table",
                    "how_to_fix": "Исправь дату на вид '06.12.2024' или '06.12.2024 г.' и проверь день/месяц/год.",
                })
    
    if len(dates_found) >= 2:
        # Проверяем порядок
        sorted_dates = sorted(dates_found, key=lambda x: x[0])
        if dates_found != sorted_dates:
            issues.append({
                "severity": sev,
                "type": "logical",
                "rule": "CalendarPlan.Order",
                "message": "Нарушен хронологический порядок дат в календарном плане.",
                "evidence": f"Найдено дат: {len(dates_found)}",
                "location": "calendar_plan_table",
                "how_to_fix": "Упорядочь даты этапов работы в хронологическом порядке (от ранних к поздним).",
            })
    
    return issues


def rule_formatting_docx(dm: DocumentModel, profile: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Упрощённая проверка форматирования DOCX: оставляем только проверки
    полужирного в теле (если запрещено) и курсивных фрагментов без латиницы.
    """
    cfg = profile.get("formatting_docx") or {}
    if not cfg.get("enabled", False):
        return []
    if dm.fmt != "docx":
        return []

    sev = cfg.get("severity", "warning")
    bold_cfg = (cfg.get("bold") or {})
    bold_only_headings = bool(bold_cfg.get("allow_in_headings_only", True))

    italic_cfg = (cfg.get("italic") or {})
    warn_not_latin = bool(italic_cfg.get("warn_if_not_latin", True))

    doc = dm.meta.get("_docx")
    if doc is None:
        return []

    issues: List[Dict[str, Any]] = []
    italic_non_latin = 0
    samples_italic = []

    for pi, p in enumerate(doc.paragraphs[:900]):
        ptext = (p.text or "").strip()
        if not ptext:
            continue

        style = (p.style.name or "").lower() if p.style is not None else ""

        for r in p.runs:
            rt = (r.text or "").strip()
            if not rt:
                continue

            if warn_not_latin and is_italic(r):
                if not is_latin_text(rt):
                    italic_non_latin += 1
                    if len(samples_italic) < 5:
                        samples_italic.append((pi, rt[:40]))

    if italic_non_latin > 0:
        issues.append({
            "severity": "info",
            "type": "formal",
            "rule": "Font.ItalicNonLatin",
            "message": f"Курсив обычно используют для латинских терминов (in vivo/in vitro и т.п.). Найдено курсивных фрагментов без латиницы: {italic_non_latin}.",
            "evidence": "; ".join([f"p{pi}: '{txt}'" for pi, txt in samples_italic]),
            "location": "docx:runs",
            "how_to_fix": "Проверь, что курсив использован по правилам методички (латынь/термины).",
        })

    return issues


def extract_table_of_contents(text: str) -> List[str]:
    """
    Извлекает список разделов из оглавления/содержания.
    Возвращает список названий разделов.
    """
    sections = []
    text_low = text.lower()

    # Ищем начало оглавления
    toc_markers = ["содержание", "оглавление"]
    toc_start = -1
    for marker in toc_markers:
        pos = text_low.find(marker)
        if pos != -1:
            toc_start = pos
            break

    if toc_start == -1:
        return []

    # Ищем конец оглавления (обычно до "введение" или следующего большого раздела)
    toc_end_markers = ["введение", "1.", "глава 1", "раздел 1"]
    toc_end = len(text)
    for marker in toc_end_markers:
        pos = text_low.find(marker, toc_start + 100)
        if pos != -1 and pos < toc_end:
            toc_end = pos

    toc_text = text[toc_start:toc_end]

    # Извлекаем строки, похожие на разделы (с номерами или заглавными буквами)
    lines = toc_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # Пропускаем заголовки самого оглавления
        if any(m in line.lower() for m in toc_markers):
            continue
        # Ищем строки с номерами разделов или заглавными буквами
        if re.match(r'^\d+[\.\)]\s+', line) or re.match(r'^[А-ЯЁ]\.\s+', line):
            # Убираем номера и точки
            clean_line = re.sub(r'^\d+[\.\)]\s+', '', line)
            clean_line = re.sub(r'^[А-ЯЁ]\.\s+', '', clean_line)
            # Убираем номера страниц в конце
            clean_line = re.sub(r'\s+\d+\s*$', '', clean_line)
            if clean_line and len(clean_line) > 3:
                sections.append(clean_line.strip())

    return sections


def rule_content_match(dm: DocumentModel, profile: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Проверка соответствия содержания пояснительной записки фактическому плану работы.
    Сравнивает разделы из оглавления с реальными разделами в документе.
    """
    cfg = profile.get("content_match") or {}
    if not cfg.get("enabled", False):
        return []

    sev = cfg.get("severity", "critical")
    issues: List[Dict[str, Any]] = []

    # Извлекаем оглавление
    toc_sections = extract_table_of_contents(dm.text)

    if not toc_sections:
        return []

    # Извлекаем реальные разделы из документа (заголовки)
    actual_sections = [h.text for h in dm.headings if h.level <= 2]

    if not actual_sections:
        issues.append({
            "severity": "warning",
            "type": "structural",
            "rule": "ContentMatch.NoHeadings",
            "message": "Не найдены заголовки разделов в документе.",
            "evidence": "—",
            "location": "document",
            "how_to_fix": "Проверь, что разделы оформлены как заголовки (стили Heading в Word).",
        })
        return issues

    # Нормализуем для сравнения
    def normalize_section(s: str) -> str:
        s = re.sub(r'\s+', ' ', s.lower().strip())
        # Убираем номера разделов
        s = re.sub(r'^\d+[\.\)]\s+', '', s)
        s = re.sub(r'^[А-ЯЁ]\.\s+', '', s)
        return s.strip()

    toc_normalized = [normalize_section(s) for s in toc_sections]
    actual_normalized = [normalize_section(s) for s in actual_sections]

    # Проверяем, все ли разделы из оглавления есть в документе
    missing_in_doc = []
    for toc_sec in toc_normalized:
        if not any(toc_sec in act or act in toc_sec for act in actual_normalized):
            missing_in_doc.append(toc_sec)

    if missing_in_doc:
        issues.append({
            "severity": sev,
            "type": "structural",
            "rule": "ContentMatch.MissingSections",
            "message": f"Разделы из оглавления отсутствуют в документе: {', '.join(missing_in_doc[:5])}",
            "evidence": f"Найдено в оглавлении: {len(toc_sections)}, найдено в документе: {len(actual_sections)}",
            "location": "document",
            "how_to_fix": "Проверь соответствие оглавления фактическому содержанию работы. Добавь отсутствующие разделы или исправь оглавление.",
        })

    # Проверяем порядок разделов (базовая проверка)
    if len(toc_normalized) >= 2 and len(actual_normalized) >= 2:
        matches = 0
        for i in range(min(3, len(toc_normalized), len(actual_normalized))):
            if toc_normalized[i] in actual_normalized[i] or actual_normalized[i] in toc_normalized[i]:
                matches += 1

        if matches == 0:
            issues.append({
                "severity": "warning",
                "type": "structural",
                "rule": "ContentMatch.Order",
                "message": "Возможно нарушен порядок разделов: оглавление не соответствует фактическому порядку в документе.",
                "evidence": f"Первые разделы в оглавлении: {', '.join(toc_normalized[:3])}",
                "location": "table_of_contents",
                "how_to_fix": "Проверь порядок разделов в документе и обнови оглавление.",
            })

    return issues


def rule_topic_match(dm: DocumentModel, profile: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Проверка соответствия темы работы утверждённому образцу.
    """
    cfg = profile.get("topic_match") or {}
    if not cfg.get("enabled", False):
        return []

    sev = cfg.get("severity", "critical")
    issues: List[Dict[str, Any]] = []

    approved_topics = cfg.get("approved_topics", [])
    if not approved_topics:
        return []  # Нет списка утверждённых тем

    title_scope = dm.text[:6000] if len(dm.text) > 6000 else dm.text
    text_low = title_scope.lower()

    # Ищем тему работы
    topic_patterns = ["тема", "на тему", "тема работы", "тема вкр"]
    found_topic = None

    for pattern in topic_patterns:
        if pattern in text_low:
            pos = text_low.find(pattern)
            # Берём текст после маркера темы
            after_marker = title_scope[pos:pos+300]
            # Извлекаем тему (до следующего заголовка или конца)
            topic_match = re.search(r"тема[:\s]+(.+?)(?:\n|$|руководитель|год|город)", after_marker, re.IGNORECASE)
            if topic_match:
                found_topic = topic_match.group(1).strip()
                break

    if not found_topic:
        issues.append({
            "severity": "warning",
            "type": "formal",
            "rule": "Topic.NotFound",
            "message": "Не найдена тема работы в титульной зоне.",
            "evidence": "—",
            "location": "title_page",
            "how_to_fix": "Проверь, что тема работы указана на титульном листе.",
        })
        return issues

    found_topic_norm = re.sub(r"\s+", " ", found_topic.lower().strip())

    matches = False
    for approved in approved_topics:
        approved_norm = re.sub(r"\s+", " ", approved.lower().strip())
        if found_topic_norm == approved_norm or found_topic_norm in approved_norm or approved_norm in found_topic_norm:
            matches = True
            break

    if not matches:
        issues.append({
            "severity": sev,
            "type": "formal",
            "rule": "Topic.Mismatch",
            "message": f"Тема работы не соответствует утверждённому образцу. Найдено: '{found_topic[:100]}'",
            "evidence": found_topic[:150],
            "location": "title_page",
            "how_to_fix": "Проверь формулировку темы работы и сверь с утверждённым вариантом.",
        })

    return issues

# улучшенный regex для строгого числового формата ДД.ММ.ГГГГ с учётом неразрывного пробела и вариантов "г."
STRICT_NUMERIC_DATE_RE = re.compile(r'\b\d{2}\.\d{2}\.\d{4}(?:[\u00A0\s]*г\.)?\b', flags=re.IGNORECASE)

def is_strict_numeric_date(text: str) -> bool:
    return bool(STRICT_NUMERIC_DATE_RE.search(text))

def run_all_rules(dm: DocumentModel, profile: Dict[str, Any]) -> Dict[str, Any]:
    issues: List[Dict[str, Any]] = []
    issues += rule_title_fields(dm, profile)
    issues += rule_dates(dm, profile)
    issues += rule_fio_cases(dm, profile)
    

    issues += rule_student_fio_detailed(dm, profile)
    issues += rule_supervisor_fio_detailed(dm, profile)
    issues += rule_date_format_strict(dm, profile)
    issues += rule_text_formatting(dm, profile)
    issues += rule_calendar_plan(dm, profile)
    issues += rule_topic_match(dm, profile)
    issues += rule_content_match(dm, profile)

    issues += rule_margins_docx(dm, profile)
    issues += rule_indent_docx(dm, profile)
    # formatting checks for DOCX (kept minimal: bold in body and italic without latin)
    issues += rule_formatting_docx(dm, profile)

    def cnt(s: str) -> int:
        return sum(1 for it in issues if (it.get("severity") or "").lower() == s)

    summary = {
        "critical": cnt("critical"),
        "warning": cnt("warning"),
        "info": cnt("info"),
        "total": len(issues),
    }

    # не отдаём _docx наружу
    public_meta = {k: v for k, v in (dm.meta or {}).items() if not str(k).startswith("_")}
    detected = (dm.meta or {}).get("detected", {})
    return {
        "profile": profile.get("name", "profile"),
        "format": dm.fmt,
        "pages": dm.pages,
        "meta": public_meta,
        "detected": {
        **detected,
        "headings_found": len(dm.headings),
        "sample_headings": [{"text": h.text[:80], "location": h.location} for h in dm.headings[:10]],
    },
        "summary": summary,
        "issues": issues,
    }


# ----------------------- API -----------------------

@app.get("/api/health")
def health():
    return {"ok": True}


@app.post("/api/check")
async def check(file: UploadFile = File(...), profile: str = "vkr_ru"):
    name = (file.filename or "").lower().strip()
    if not (name.endswith(".docx") or name.endswith(".pdf")):
        raise HTTPException(status_code=400, detail="Поддерживаются только .docx и .pdf (цифровой PDF без OCR)")

    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Пустой файл")

    prof = load_profile(profile)

    try:
        if name.endswith(".docx"):
            dm = extract_docx(data)
        else:
            dm = extract_pdf(data)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Не удалось извлечь текст/структуру: {e}")

    report = run_all_rules(dm, prof)
    return report
